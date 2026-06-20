#!/usr/bin/env python3
"""
aikit/tools/convert_requirements.py

Cross-platform requirement document converter for aikit.
Converts .pdf, .xlsx, .xls, .docx, .doc, .csv, .txt, .md -> Markdown.

Works on macOS, Windows, and Linux. Requires Python 3.9+.

Usage:
    python convert_requirements.py <input_dir> [output_dir]

    input_dir   Folder containing requirement documents for a module.
                Pipeline will NOT start if this folder is empty.
    output_dir  (optional) Destination for converted .md files.
                Defaults to <input_dir>/converted/

Install dependencies first:
    pip install -r aikit/tools/requirements.txt

Examples:
    python aikit/tools/convert_requirements.py aikit/inputs/requirements/ecommerce
    python aikit/tools/convert_requirements.py aikit/inputs/requirements/ecommerce aikit/outputs/ecommerce/requirements/converted
"""

import csv
import json
import sys
from datetime import datetime
from pathlib import Path

SUPPORTED_EXTENSIONS = {".pdf", ".xlsx", ".xls", ".docx", ".doc", ".csv", ".txt", ".md"}


# ── Converters ────────────────────────────────────────────────────────────────


def pdf_to_markdown(file_path: Path) -> str:
    try:
        import pdfplumber
    except ImportError:
        return (
            f"# {file_path.stem}\n\n"
            f"> **Error**: `pdfplumber` is not installed.\n"
            f"> Run: `pip install pdfplumber`\n"
        )

    lines = [f"# {file_path.stem}\n\n*Source: {file_path.name}*\n"]
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []

            if text.strip():
                lines.append(f"\n## Page {page_num}\n\n{text.strip()}\n")

            for table in tables:
                if not table:
                    continue
                header = [str(c or "").strip() for c in table[0]]
                lines.append("\n### Table\n")
                lines.append("| " + " | ".join(header) + " |")
                lines.append("| " + " | ".join("---" for _ in header) + " |")
                for row in table[1:]:
                    cells = [str(c or "").strip() for c in row]
                    # Pad row to header length
                    while len(cells) < len(header):
                        cells.append("")
                    lines.append("| " + " | ".join(cells[: len(header)]) + " |")
                lines.append("")

    return "\n".join(lines)


def xlsx_to_markdown(file_path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        return (
            f"# {file_path.stem}\n\n"
            f"> **Error**: `openpyxl` is not installed.\n"
            f"> Run: `pip install openpyxl`\n"
        )

    lines = [f"# {file_path.stem}\n\n*Source: {file_path.name}*\n"]
    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    except Exception as exc:
        lines.append(f"> **Error loading workbook**: {exc}\n")
        return "\n".join(lines)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = [
            [str(c) if c is not None else "" for c in row]
            for row in ws.iter_rows(values_only=True)
        ]
        # Drop completely empty rows
        rows = [r for r in rows if any(c.strip() for c in r)]
        if not rows:
            continue

        lines.append(f"\n## Sheet: {sheet_name}\n")
        header = rows[0]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        for row in rows[1:]:
            while len(row) < len(header):
                row.append("")
            lines.append("| " + " | ".join(row[: len(header)]) + " |")
        lines.append("")

    return "\n".join(lines)


def xls_to_markdown(file_path: Path) -> str:
    try:
        import xlrd
    except ImportError:
        return (
            f"# {file_path.stem}\n\n"
            f"> **Error**: `xlrd` is not installed for legacy .xls support.\n"
            f"> Run: `pip install xlrd`\n"
        )

    lines = [f"# {file_path.stem}\n\n*Source: {file_path.name}*\n"]
    try:
        wb = xlrd.open_workbook(str(file_path))
    except Exception as exc:
        lines.append(f"> **Error loading workbook**: {exc}\n")
        return "\n".join(lines)

    for sheet in wb.sheets():
        rows = [
            [str(sheet.cell_value(r, c)) for c in range(sheet.ncols)]
            for r in range(sheet.nrows)
        ]
        rows = [r for r in rows if any(c.strip() for c in r)]
        if not rows:
            continue
        lines.append(f"\n## Sheet: {sheet.name}\n")
        header = rows[0]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")
        for row in rows[1:]:
            lines.append("| " + " | ".join(row) + " |")
        lines.append("")

    return "\n".join(lines)


def docx_to_markdown(file_path: Path) -> str:
    # Prefer mammoth — produces cleaner markdown from Word styles
    try:
        import mammoth

        with open(file_path, "rb") as fh:
            result = mammoth.convert_to_markdown(fh)
        md = f"# {file_path.stem}\n\n*Source: {file_path.name}*\n\n{result.value}"
        if result.messages:
            notes = "\n".join(f"- {m}" for m in result.messages)
            md += f"\n\n> **Conversion notes:**\n{notes}"
        return md
    except ImportError:
        pass  # fall through to python-docx
    except Exception:
        pass  # malformed file — fall through

    # Fallback: python-docx
    try:
        from docx import Document  # type: ignore

        doc = Document(str(file_path))
        lines = [f"# {file_path.stem}\n\n*Source: {file_path.name}*\n"]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            style_name = para.style.name if para.style else ""
            if "Heading 1" in style_name:
                lines.append(f"\n# {text}\n")
            elif "Heading 2" in style_name:
                lines.append(f"\n## {text}\n")
            elif "Heading 3" in style_name:
                lines.append(f"\n### {text}\n")
            else:
                lines.append(text)

        for table in doc.tables:
            lines.append("")
            for i, row in enumerate(table.rows):
                cells = [c.text.strip().replace("\n", " ") for c in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    lines.append("| " + " | ".join("---" for _ in cells) + " |")
            lines.append("")

        return "\n".join(lines)

    except ImportError:
        return (
            f"# {file_path.stem}\n\n"
            f"> **Error**: Neither `mammoth` nor `python-docx` is installed.\n"
            f"> Run: `pip install mammoth python-docx`\n"
        )
    except Exception as exc:
        return (
            f"# {file_path.stem}\n\n"
            f"> **Error opening .docx**: {exc}\n"
            f"> Ensure the file is a valid .docx format (not password-protected).\n"
        )


def doc_to_markdown(file_path: Path) -> str:
    """
    Legacy .doc (OLE2 binary) converter.
    Tries python-docx first (works when the file is actually DOCX with a .doc extension).
    Otherwise returns a clear user-facing message with conversion steps.
    """
    # Some files are mislabelled — try as DOCX first
    result = docx_to_markdown(file_path)
    if "> **Error" not in result:
        return result

    return (
        f"# {file_path.stem}\n\n"
        f"*Source: {file_path.name}*\n\n"
        f"> **Warning**: Legacy `.doc` format (OLE2 binary) cannot be read directly.\n"
        f"> Please convert this file to `.docx` using one of these methods:\n"
        f"> \n"
        f"> **Windows (PowerShell):**\n"
        f"> ```powershell\n"
        f"> $word = New-Object -ComObject Word.Application\n"
        f"> $doc  = $word.Documents.Open(\"$PWD\\{file_path.name}\")\n"
        f"> $doc.SaveAs([ref]\"$PWD\\{file_path.stem}.docx\", [ref]16)\n"
        f"> $doc.Close(); $word.Quit()\n"
        f"> ```\n"
        f"> \n"
        f"> **macOS / Linux (requires LibreOffice):**\n"
        f"> ```bash\n"
        f"> soffice --headless --convert-to docx \"{file_path.name}\"\n"
        f"> ```\n"
        f"> \n"
        f"> **Online**: Upload to Google Docs → Download as .docx\n"
    )


def csv_to_markdown(file_path: Path) -> str:
    lines = [f"# {file_path.stem}\n\n*Source: {file_path.name}*\n"]
    try:
        # utf-8-sig strips BOM written by Excel on Windows
        with open(file_path, encoding="utf-8-sig", newline="") as fh:
            reader = csv.reader(fh)
            rows = [r for r in reader if any(c.strip() for c in r)]
    except UnicodeDecodeError:
        with open(file_path, encoding="latin-1", newline="") as fh:
            reader = csv.reader(fh)
            rows = [r for r in reader if any(c.strip() for c in r)]

    if not rows:
        return f"# {file_path.stem}\n\n*Empty file.*\n"

    header = rows[0]
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in rows[1:]:
        while len(row) < len(header):
            row.append("")
        lines.append("| " + " | ".join(row[: len(header)]) + " |")
    lines.append("")
    return "\n".join(lines)


def txt_to_markdown(file_path: Path) -> str:
    try:
        content = file_path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError:
        content = file_path.read_text(encoding="latin-1")
    return f"# {file_path.stem}\n\n*Source: {file_path.name}*\n\n{content}\n"


def md_passthrough(file_path: Path) -> str:
    try:
        return file_path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError:
        return file_path.read_text(encoding="latin-1")


CONVERTERS = {
    ".pdf": pdf_to_markdown,
    ".xlsx": xlsx_to_markdown,
    ".xls": xls_to_markdown,
    ".docx": docx_to_markdown,
    ".doc": doc_to_markdown,
    ".csv": csv_to_markdown,
    ".txt": txt_to_markdown,
    ".md": md_passthrough,
}


# ── Orchestrator ──────────────────────────────────────────────────────────────


def convert_directory(input_dir: Path, output_dir: Path) -> dict:
    output_dir.mkdir(parents=True, exist_ok=True)

    candidates = sorted(
        f for f in input_dir.iterdir()
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    )

    if not candidates:
        print(f"\n[aikit] ERROR: No supported files found in: {input_dir}")
        print(f"  Supported formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}")
        print("  Add your requirement documents to this folder before running aikit.")
        sys.exit(1)

    print(f"\n[aikit] Found {len(candidates)} file(s) in {input_dir}")

    results: dict = {"converted": [], "skipped": [], "failed": [], "errors": []}

    for file_path in candidates:
        ext = file_path.suffix.lower()
        converter = CONVERTERS.get(ext)
        if converter is None:
            results["skipped"].append(file_path.name)
            print(f"  [SKIP]    {file_path.name}")
            continue

        out_path = output_dir / f"{file_path.stem}.md"
        print(f"  [CONVERT] {file_path.name} -> {out_path.name} ... ", end="", flush=True)
        try:
            markdown = converter(file_path)
            out_path.write_text(markdown, encoding="utf-8")
            results["converted"].append(file_path.name)
            print("OK")
        except Exception as exc:
            results["failed"].append(file_path.name)
            results["errors"].append(f"{file_path.name}: {exc}")
            print(f"FAILED ({exc})")

    # Write manifest for the pipeline to read
    manifest = {
        "generated_at": datetime.now().isoformat(),
        "input_dir": str(input_dir),
        "output_dir": str(output_dir),
        "results": results,
    }
    manifest_path = output_dir / "_manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    print(f"\n[aikit] Conversion complete:")
    print(f"  Converted : {len(results['converted'])}")
    print(f"  Skipped   : {len(results['skipped'])}")
    print(f"  Failed    : {len(results['failed'])}")
    print(f"  Output    : {output_dir}")

    if results["failed"]:
        print("\n[aikit] WARNING — failed files:")
        for err in results["errors"]:
            print(f"  - {err}")

    if not results["converted"]:
        print("\n[aikit] ERROR: No files were successfully converted. Pipeline cannot start.")
        sys.exit(1)

    return results


def main() -> None:
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    input_dir = Path(sys.argv[1]).resolve()

    if not input_dir.exists():
        print(f"[aikit] ERROR: Input directory not found: {input_dir}")
        sys.exit(1)
    if not input_dir.is_dir():
        print(f"[aikit] ERROR: Not a directory: {input_dir}")
        sys.exit(1)

    output_dir = Path(sys.argv[2]).resolve() if len(sys.argv) >= 3 else input_dir / "converted"

    convert_directory(input_dir, output_dir)


if __name__ == "__main__":
    main()
